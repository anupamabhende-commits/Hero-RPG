from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"E:\Hero-RPG\Player_Code_Explained.docx"


def set_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = tcPr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcPr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)


def para(doc, text="", style=None, before=0, after=6, line=1.1, color=None, size=11, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), size=11)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color="2E74B5" if level == 1 else "1F4D78")
    return p


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
styles["Normal"].font.size = Pt(11)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.paragraph_format.space_after = Pt(0)
set_font(header.add_run("HERO RPG | PLAYER CODE GUIDE"), size=8, bold=True, color="6B7280")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Simple explanation of the current player scripts"), size=8, color="6B7280")

# Title
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Player Code Explained")
set_font(r, size=26, bold=True, color="17365D")
para(doc, "A simple guide to how the player moves, fights, casts spells, and uses portals in Hero RPG.", after=14, size=12, color="4B5563")

intro = doc.add_table(rows=1, cols=1)
intro.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_widths(intro, [6.5])
cell = intro.cell(0, 0)
shade(cell, "EAF2F8")
cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
lead = p.add_run("Big idea: ")
set_font(lead, bold=True, color="17365D")
set_font(p.add_run("the Player object is not controlled by one huge script. It is a team of small scripts, where each script has one clear job."), color="17365D")

heading(doc, "1. The player is made of small systems")
para(doc, "The Player GameObject contains several components. They share information through PlayerController, which acts like the coordinator.")
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_widths(table, [2.1, 4.4])
for i, label in enumerate(["Script", "Simple job"]):
    c = table.rows[0].cells[i]
    shade(c, "E8EEF5")
    set_font(c.paragraphs[0].add_run(label), bold=True, color="17365D")
rows = [
    ("PlayerInput", "Reads the keyboard and mouse. It remembers what the player pressed this frame."),
    ("PlayerController", "The main coordinator. It asks the other systems to move, rotate, jump, fight, or cast."),
    ("PlayerMovement", "Moves the CharacterController using the current input and movement settings."),
    ("PlayerRotation", "Turns the character toward the desired direction, including lock-on behavior."),
    ("PlayerJump + GroundChecker", "Checks whether the player is on the ground, then applies jumping and gravity."),
    ("PlayerCombat", "Receives attack/magic/portal requests and starts the right action."),
    ("PlayerAnimator", "Sends values such as Speed, Grounded, Sprint, and triggers to the Animator."),
    ("PlayerSpellController", "Creates the fireball and runtime portal. It also manages casting cooldown time."),
    ("ActionRPGCamera", "Follows the player and handles camera angle/collision behavior."),
]
for script, job in rows:
    cells = table.add_row().cells
    for cell in cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_font(cells[0].paragraphs[0].add_run(script), bold=True, color="1F4D78")
    set_font(cells[1].paragraphs[0].add_run(job))

heading(doc, "2. What happens every frame")
para(doc, "Unity calls Update() once every frame. In this project, the basic order of events is:")
for item in [
    "PlayerInput reads keys such as WASD, Space, Q, F, Tab, mouse buttons, and Shift.",
    "PlayerController reads that input and tells movement, rotation, combat, and animation systems what to do.",
    "PlayerMovement changes the CharacterController position.",
    "PlayerAnimator updates the animation parameters so the model looks like it is walking, jumping, rolling, or casting.",
]: bullet(doc, item)
add_code(doc, "Keyboard/mouse -> PlayerInput -> PlayerController -> specialist script -> CharacterController / Animator")

heading(doc, "3. Movement, jumping, and camera")
heading(doc, "Movement", 2)
para(doc, "PlayerInput stores a movement direction. PlayerMovement uses the walk speed, sprint speed, acceleration, and deceleration from PlayerStats. The CharacterController is then moved safely, so it can collide with the ground and walls.")
heading(doc, "Jumping", 2)
para(doc, "GroundChecker checks whether the CharacterController is standing on ground. When the jump key is pressed and the player is grounded, PlayerJump adds upward movement. Gravity then pulls the player back down.")
heading(doc, "Camera and facing direction", 2)
para(doc, "ActionRPGCamera follows a camera target near the player. PlayerRotation uses the camera direction so pressing forward means “move where the camera is looking,” not simply “move along the world’s Z axis.”")

heading(doc, "4. Combat and spells")
para(doc, "PlayerCombat is the bridge between the normal player controls and the magic system. It checks whether the player is busy, then asks PlayerSpellController to begin a spell.")
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_widths(table, [1.4, 5.1])
for i, label in enumerate(["Input", "What the code does"]):
    shade(table.rows[0].cells[i], "E8EEF5")
    set_font(table.rows[0].cells[i].paragraphs[0].add_run(label), bold=True, color="17365D")
for key, action in [
    ("Left mouse", "PlayerCombat starts the normal attack animation."),
    ("Right mouse", "PlayerCombat starts the magic/cast action."),
    ("1", "PlayerSpellController starts the fireball spell directly."),
    ("F", "PlayerCombat asks PlayerSpellController to cast a portal. The spell script also currently listens for F itself."),
]:
    cells = table.add_row().cells
    set_font(cells[0].paragraphs[0].add_run(key), bold=True, color="1F4D78")
    set_font(cells[1].paragraphs[0].add_run(action))

heading(doc, "5. How a fireball works")
para(doc, "The fireball is connected to the casting animation. Animation events can call methods in PlayerSpellController at the correct moment.")
for item in [
    "CastFireball() checks whether another spell is already active.",
    "The Animator receives the Fireball trigger, starting the cast animation.",
    "SpawnFireball() creates a copy of the fireball prefab at the fire point.",
    "EnableFireballVFX() can show the charging effect while the spell is being prepared.",
    "DisableFireballVFX() launches it. The Fireball script moves it forward until it reaches maxDistance, then destroys it.",
]: bullet(doc, item)
add_code(doc, "PlayerSpellController -> Animator event -> SpawnFireball -> Fireball.Launch(direction) -> move -> destroy")

heading(doc, "6. How the portal works")
para(doc, "Pressing F starts CastPortal(). After a short delay, the code creates a new GameObject called Runtime Portal in front of the player. PortalTeleportController builds the orange ring, blue interior, and sparks at runtime.")
para(doc, "When the player crosses the portal opening, TeleportPlayer() moves the Player transform to the Portal Destination object in the scene. It temporarily disables the CharacterController first, which avoids collision problems while changing position.")
heading(doc, "Important portal settings", 2)
for item in [
    "portalSpawnDelay: how long the game waits before the portal appears.",
    "portalForwardOffset: how far in front of the player the portal is created.",
    "portalRadius: the size of the portal’s usable opening.",
    "portalLifetime: how long it stays open.",
    "portalDestination: the Transform where the player arrives.",
]: bullet(doc, item)
para(doc, "Note: the current portal check measures the player's horizontal and vertical distance from the portal center. Since the portal is raised above ground, use only horizontal distance for the entry test or lower the portal's vertical offset. Otherwise walking through the visual opening may not teleport the player.", before=4, after=8, color="7A5A00", bold=True)

heading(doc, "7. Easy way to explain the code in a demo")
for item in [
    "“Input reads what I press.”",
    "“The controller decides which player system should react.”",
    "“Movement, jumping, combat, animation, and spells are separate so each part is easier to change.”",
    "“The spell controller creates the fireball or portal, and the portal controller handles the actual teleport.”",
    "“PlayerStats keeps values such as speed and gravity in one place, so tuning the feel of the game is simple.”",
]: bullet(doc, item)

heading(doc, "8. Files to open when explaining")
for item in [
    "Assets/Player/Scripts/PlayerInput.cs - where keyboard and mouse actions are read.",
    "Assets/Player/Scripts/PlayerController.cs - the player coordinator.",
    "Assets/Player/Scripts/PlayerMovement.cs and PlayerRotation.cs - moving and facing.",
    "Assets/Player/Scripts/PlayerCombat.cs - combat decisions.",
    "Assets/Scripts/PlayerSpellController.cs - fireball and portal casting.",
    "Assets/Scripts/PortalTeleportController.cs - portal visuals, detection, and teleporting.",
]: bullet(doc, item)

doc.save(OUT)
print(OUT)
